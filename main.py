#!/usr/bin/env python3
# ==============================================================================
# main.py -- Storage QC daily report (form 573399)
#
# Pulls a day's photos, scores them with the reason models, and produces:
#     reports/storage_<date>.html   full report with thumbnails
#     reports/storage_<date>.pdf    same, hyperlinks clickable
#     reports/storage_<date>.docx   Word version
#     reports/storage_<date>.csv    flat data for tracking repeats
# then emails the Arabic summary table as the message body with the PDF and
# Word attached.
#
# Repo layout:
#     models/field_9199839/reason_models_9199839.joblib
#     models/field_9199840/reason_models_9199840.joblib
#     models/field_9199841/reason_models_9199841.joblib
#
# Only two secrets: ZENPUT_TOKEN and EMAIL_PASSWORD.
# ==============================================================================

import os, ast, csv, ssl, base64, hashlib, smtplib, warnings, re
import datetime as dt
from concurrent.futures import ThreadPoolExecutor
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.utils import formataddr
from io import BytesIO
from collections import defaultdict

import numpy as np
import requests, pytz
from requests.adapters import HTTPAdapter
from PIL import Image, ImageOps

warnings.filterwarnings("ignore")

# ------------------------------- EMAIL ----------------------------------------
SENDER_EMAIL   = "aof.group.auto@gmail.com"
SENDER_NAME    = "Business Intelligence"
SENDER_PASS    = os.environ.get("EMAIL_PASSWORD", "")     # app password, secret
SMTP_HOST      = "smtp.gmail.com"
SMTP_PORT      = 587

RECIPIENTS     = ["o.salahaddin@aofgroup.com"]
CC             = ["a.alsalem@aofgroup.com"]

# ------------------------------- CONFIG ---------------------------------------
API_TOKEN   = os.environ.get("ZENPUT_TOKEN", "")
TEMPLATE_ID = 573399
DAY_OFFSET  = int(os.environ.get("DAY_OFFSET", "1"))     # 1 = yesterday
REPORT_DATE = os.environ.get("REPORT_DATE", "")          # or "20260830"

FIELDS = {
    9199839: ("main chiller", "Main Chiller"),
    9199840: ("main freezer", "Freezer"),
    9199841: ("dry store",    "Dry Store"),
}

PRIORITY    = ["overstacking", "arrangement", "boxcut", "wrongphoto"]
PHOTO_ISSUE = "wrongphoto"

LABELS = {
    "overstacking": "Overstacking — boxes leaning, overhanging or above the rail",
    "arrangement":  "Arrangement — items off the shelves or disordered",
    "wrongphoto":   "Wrong or unusable photo — storage area not properly shown",
    "boxcut":       "Cut or torn boxes",
}

ARABIC = {
    "overstacking": "تكدس",
    "arrangement":  "عدم تنظيم",
    "boxcut":       "كراتين ممزقة",
    "wrongphoto":   "صور غير صحيحة",
}

COLOURS = {
    9199839: ("#BDD7EE", "#9DC3E6"),   # chiller
    9199840: ("#FF5B5B", "#FF3B3B"),   # freezer
    9199841: ("#FBE2D5", "#F4B183"),   # dry store
}
NOTE_BG = "#FF0000"

RED_AT     = 50
MODEL_NAME = "facebook/dinov2-base"
MODEL_DIR  = "models"
OUT_DIR    = "reports"
TMP        = "photos"
BATCH, WORKERS = 32, 8

TZ      = pytz.timezone("Asia/Riyadh")
BASE    = "https://www.zenput.com"
STORAGE = f"{BASE}/api/v2/users/current/storage/"
HEADERS = {"X-API-TOKEN": API_TOKEN}
NOW     = dt.datetime.now(TZ)
DAY     = REPORT_DATE or (NOW - dt.timedelta(days=DAY_OFFSET)).strftime("%Y%m%d")
PRETTY  = f"{DAY[:4]}-{DAY[4:6]}-{DAY[6:]}"

if not API_TOKEN:
    raise SystemExit("ZENPUT_TOKEN not set")
os.makedirs(TMP, exist_ok=True)
os.makedirs(OUT_DIR, exist_ok=True)
print(f"form {TEMPLATE_ID} | reporting on {DAY}")


# ------------------------------ MODELS ----------------------------------------
import torch, joblib
from transformers import AutoImageProcessor, AutoModel

BUNDLES = {}
for fid, (fname, _) in FIELDS.items():
    p = f"{MODEL_DIR}/reason_models_{fid}.joblib"
    if not os.path.exists(p):
        print(f"  !! missing {p}")
        continue
    BUNDLES[fid] = joblib.load(p)
    print(f"  {fname}: {', '.join(BUNDLES[fid]['reasons'])}")
if not BUNDLES:
    raise SystemExit("No models found")


def to_risk(p, thr):
    """Rescale so each model's own threshold lands on 50%."""
    if thr <= 0 or thr >= 1:
        return 100.0 * p
    return 50.0 * p / thr if p < thr else 50.0 + 50.0 * (p - thr) / (1.0 - thr)


# ------------------------------ ZENPUT ----------------------------------------
sess = requests.Session()
_ad = HTTPAdapter(pool_connections=WORKERS * 2, pool_maxsize=WORKERS * 2, max_retries=2)
sess.mount("https://", _ad); sess.mount("http://", _ad)


def sub_date(s):
    md = s.get("smetadata") or {}
    for src in (md, s):
        for k in ("date_submitted", "submitted_at", "date_created", "created_at"):
            v = src.get(k)
            if v:
                try:
                    d = dt.datetime.fromisoformat(str(v).replace("Z", "+00:00"))
                    return d if d.tzinfo else TZ.localize(d)
                except Exception:
                    pass
    return None


def sub_branch(s):
    loc = (s.get("smetadata") or {}).get("location") or {}
    return str(loc.get("name") or loc.get("code") or loc.get("id") or "unknown") \
        if isinstance(loc, dict) else str(loc or "unknown")


def branch_code(name):
    m = re.search(r"\bB\s?(\d{1,3})\b", name, re.I)
    return f"B{int(m.group(1)):02d}" if m else name[:8]


def sub_user(s):
    u = (s.get("smetadata") or {}).get("created_by") or {}
    return str(u.get("display_name", "")) if isinstance(u, dict) else ""


def flat(s):
    out = []
    def walk(n):
        if isinstance(n, list):
            for x in n: walk(x)
        elif isinstance(n, dict):
            if "field_type" in n or "title" in n: out.append(n)
            for k in ("answers", "children", "sub_answers", "items"):
                if k in n: walk(n[k])
    walk(s.get("answers") or [])
    return out


def photo_recs(val):
    """Read `value` only -- `image_value` duplicates it."""
    found = []
    def walk(v):
        if isinstance(v, dict):
            if "s3_key" in v: found.append(v)
            else:
                for x in v.values(): walk(x)
        elif isinstance(v, list):
            for x in v: walk(x)
        elif isinstance(v, str) and v[:1] in "[{":
            try: walk(ast.literal_eval(v))
            except Exception: pass
    walk(val)
    return found


def find_url(o):
    if isinstance(o, str):
        return o if o.lower().startswith("http") else None
    if isinstance(o, dict):
        for k in ("location", "url", "signed_url"):
            v = o.get(k)
            if isinstance(v, str) and v.lower().startswith("http"): return v
        for v in o.values():
            u = find_url(v)
            if u: return u
    if isinstance(o, list):
        for v in o:
            u = find_url(v)
            if u: return u
    return None


def grab(j):
    """s3_key must be exchanged for a signed URL; direct requests 403."""
    if os.path.exists(j["path"]):
        return True
    try:
        r = sess.get(STORAGE, headers=HEADERS, params={"path": j["s3_key"]}, timeout=45)
        if r.status_code != 200:
            return False
        if r.headers.get("Content-Type", "").startswith("image/"):
            data = r.content
        else:
            u = find_url(r.json())
            if not u: return False
            r2 = sess.get(u, timeout=45)
            if r2.status_code != 200 or len(r2.content) < 1024: return False
            data = r2.content
        with open(j["path"], "wb") as f:
            f.write(data)
        return True
    except Exception:
        return False


print("pulling submissions...")
subs, start = [], 0
while start < 1500:
    r = sess.get(f"{BASE}/api/v3/submissions/", headers=HEADERS,
                 params={"form_template_id": TEMPLATE_ID, "limit": 100, "start": start},
                 timeout=60)
    if r.status_code != 200:
        print(f"  !! HTTP {r.status_code} {r.text[:200]}"); break
    b = r.json()
    if isinstance(b, dict):
        b = b.get("results") or b.get("data") or []
    if not b: break
    subs += b; start += 100
    if len(b) < 100: break

# all date filtering is client-side; the API ignores date params
jobs = []
for s in subs:
    d = sub_date(s)
    ds = d.strftime("%Y%m%d") if d else ""
    if ds != DAY:
        continue
    legacy = str(s.get("legacy_submission_id", "") or "")
    for a in flat(s):
        fid = a.get("field_id") or a.get("id")
        if fid not in BUNDLES:
            continue
        for rec in photo_recs(a.get("value")):
            k = rec.get("s3_key", "")
            if not k: continue
            h = hashlib.md5(k.encode()).hexdigest()[:10]
            br = sub_branch(s)
            jobs.append({"field_id": fid, "question": FIELDS[fid][0],
                         "s3_key": k, "branch": br, "code": branch_code(br),
                         "user": sub_user(s), "date": ds,
                         "path": f"{TMP}/{h}.jpg",
                         "link": f"https://www.zenput.com/reports/#form_id/{legacy}"})

print(f"{len(jobs)} photos on {DAY}")
if not jobs:
    raise SystemExit("Nothing submitted on that day.")

with ThreadPoolExecutor(max_workers=WORKERS) as ex:
    got = list(ex.map(grab, jobs))
jobs = [j for j, o in zip(jobs, got) if o]
print(f"downloaded {len(jobs)}")


# ----------------------------- SCORING ----------------------------------------
proc  = AutoImageProcessor.from_pretrained(MODEL_NAME)
model = AutoModel.from_pretrained(MODEL_NAME).eval()


def embed(paths):
    vecs = []
    for i in range(0, len(paths), BATCH):
        imgs = []
        for p in paths[i:i + BATCH]:
            try:
                imgs.append(ImageOps.exif_transpose(Image.open(p)).convert("RGB"))
            except Exception:
                imgs.append(Image.new("RGB", (224, 224)))
        with torch.no_grad():
            a = model(**proc(images=imgs, return_tensors="pt")).last_hidden_state[:, 0]
            b = model(**proc(images=[im.transpose(Image.FLIP_LEFT_RIGHT) for im in imgs],
                             return_tensors="pt")).last_hidden_state[:, 0]
        a = torch.nn.functional.normalize(a, dim=1)
        b = torch.nn.functional.normalize(b, dim=1)
        vecs.append(torch.nn.functional.normalize(a + b, dim=1).numpy())
        print(f"  scored {min(i + BATCH, len(paths))}/{len(paths)}")
    return np.vstack(vecs)


X = embed([j["path"] for j in jobs])

for i, j in enumerate(jobs):
    reasons = BUNDLES[j["field_id"]]["reasons"]
    risks = {r: to_risk(float(m["clf"].predict_proba(X[i:i + 1])[0, 1]), m["thr"])
             for r, m in reasons.items()}
    j["risks"] = risks
    fired = [r for r in PRIORITY if risks.get(r, 0.0) >= RED_AT]
    if fired:
        j["primary"] = fired[0]
        j["risk"] = risks[fired[0]]
        j["also"] = [r for r in fired[1:] if r != PHOTO_ISSUE]
    else:
        j["primary"] = max(risks, key=risks.get) if risks else None
        j["risk"] = risks.get(j["primary"], 0.0)
        j["also"] = []
    j["photo_note"] = (risks.get(PHOTO_ISSUE, 0.0) >= RED_AT
                       and j["primary"] != PHOTO_ISSUE)
    j["flagged"] = j["risk"] >= RED_AT
    j["prank"] = PRIORITY.index(j["primary"]) if j["primary"] in PRIORITY else 99

    parts = [ARABIC[r] for r in [j["primary"]] + j["also"] if r in ARABIC]
    if j["photo_note"]:
        parts.append(ARABIC["wrongphoto"])
    j["ar_note"] = " و".join(dict.fromkeys(parts))

flagged = [j for j in jobs if j["flagged"]]
print(f"{len(flagged)} flagged of {len(jobs)}")


# --------------------------- ARABIC TABLE -------------------------------------
def arabic_table_html():
    rows = []
    for fid, (fname, en) in FIELDS.items():
        items = [j for j in flagged if j["field_id"] == fid]
        if not items:
            continue
        per = defaultdict(list)
        for j in sorted(items, key=lambda x: (x["prank"], -x["risk"])):
            per[j["code"]].append(j["ar_note"])
        bg, hdr = COLOURS[fid]
        n, first = len(per), True
        for code, notes in per.items():
            note = " و".join(dict.fromkeys(x for s in notes for x in s.split(" و")))
            span = (f'<td rowspan="{n}" style="background:{hdr};border:1px solid #333;'
                    f'padding:6px 10px;font-weight:700;text-align:center;'
                    f'vertical-align:middle">{en}</td>') if first else ""
            rows.append(
                f'<tr>'
                f'<td style="background:{bg};border:1px solid #333;padding:5px 10px;'
                f'font-weight:700;text-align:center">{code}</td>'
                f'<td style="background:{NOTE_BG};color:#000;border:1px solid #333;'
                f'padding:5px 10px;text-align:center">{note}</td>'
                f'{span}</tr>')
            first = False

    if not rows:
        return ('<p dir="rtl" style="font:15px Tahoma,Arial;color:#27ae60">'
                'لا توجد ملاحظات على مخازن الفروع اليوم.</p>')

    return (f'<table dir="rtl" style="border-collapse:collapse;font:14px Tahoma,Arial;'
            f'direction:rtl">'
            f'<tr>'
            f'<th style="background:#D9D9D9;border:1px solid #333;padding:6px 12px">'
            f'كود الفرع</th>'
            f'<th style="background:#D9D9D9;border:1px solid #333;padding:6px 12px">'
            f'ملاحظات المراجعة</th>'
            f'<th style="background:#D9D9D9;border:1px solid #333;padding:6px 12px">'
            f'مكان التخزين</th>'
            f'</tr>{"".join(rows)}</table>')


AR_TABLE = arabic_table_html()


# ------------------------------ HTML ------------------------------------------
def thumb_b64(p, px=250):
    try:
        im = ImageOps.exif_transpose(Image.open(p)).convert("RGB")
        im.thumbnail((px, px))
        b = BytesIO(); im.save(b, "JPEG", quality=72)
        return base64.b64encode(b.getvalue()).decode()
    except Exception:
        return ""


def bar(r, v):
    col = "#c0392b" if v >= RED_AT else "#e0a030" if v >= 30 else "#bbb"
    return (f'<div style="margin:2px 0;font-size:11px;color:#555">{r} '
            f'<span style="color:{col};font-weight:600">{v:.0f}%</span></div>')


def card(j):
    col = "#c0392b"
    bars = "".join(bar(r, j["risks"][r]) for r in PRIORITY if r in j["risks"])
    note = ('<div style="margin-top:6px;padding:5px 8px;background:#fdf3e3;'
            'border-left:3px solid #e0a030;font-size:11px;color:#8a5a00">'
            'Photo also poorly framed — ask for a proper wide shot.</div>'
            ) if j["photo_note"] else ""
    return (f'<div style="border:2px solid {col};border-radius:8px;padding:10px;'
            f'margin:8px 0;page-break-inside:avoid">'
            f'<table style="width:100%"><tr>'
            f'<td style="width:260px;vertical-align:top">'
            f'<img src="data:image/jpeg;base64,{thumb_b64(j["path"])}" '
            f'style="width:250px;border-radius:5px"></td>'
            f'<td style="vertical-align:top;padding-left:14px">'
            f'<div style="font-size:20px;font-weight:700;color:{col}">'
            f'{j["risk"]:.0f}%</div>'
            f'<div style="font-weight:600">{LABELS.get(j["primary"], "")}</div>'
            f'<div style="color:#666;font-size:12px;margin-bottom:6px">'
            f'{j["branch"]} &middot; {j["user"]}</div>{bars}{note}'
            f'<div style="margin-top:6px"><a href="{j["link"]}" '
            f'style="font-size:12px;color:#0645ad">open submission in Zenput</a></div>'
            f'</td></tr></table></div>')


sections = []
for fid, (fname, en) in FIELDS.items():
    sub = [j for j in flagged if j["field_id"] == fid]
    total = sum(1 for j in jobs if j["field_id"] == fid)
    if not total:
        continue
    if not sub:
        sections.append(f'<h2 style="margin-top:26px">{en}</h2>'
                        f'<p style="color:#27ae60">No issues in {total} photos.</p>')
        continue
    per = defaultdict(list)
    for j in sub:
        per[j["branch"]].append(j)
    body = []
    for b, items in sorted(per.items(),
                           key=lambda kv: (min(i["prank"] for i in kv[1]),
                                           -max(i["risk"] for i in kv[1]))):
        top = sorted(items, key=lambda x: (x["prank"], -x["risk"]))
        body.append(f'<h3 style="margin-top:16px">{b}</h3>'
                    + "".join(card(j) for j in top))
    sections.append(f'<h2 style="margin-top:26px;border-bottom:2px solid #ddd">'
                    f'{en} — {len(sub)} of {total} flagged</h2>' + "".join(body))

HTML = (f'<html><head><meta charset="utf-8"><style>@page{{size:A4;margin:14mm}}'
        f'body{{font:13px Arial,sans-serif}}</style></head><body>'
        f'<h1 style="font-size:20px">Storage checks — {PRETTY}</h1>'
        f'<p>{len(flagged)} of {len(jobs)} photos flagged.</p>'
        f'<div style="margin:18px 0">{AR_TABLE}</div>'
        f'{"".join(sections)}</body></html>')

html_path = f"{OUT_DIR}/storage_{DAY}.html"
with open(html_path, "w", encoding="utf-8") as f:
    f.write(HTML)
print(f"wrote {html_path}")


# ------------------------------- PDF ------------------------------------------
pdf_path = f"{OUT_DIR}/storage_{DAY}.pdf"
try:
    from weasyprint import HTML as WHTML
    WHTML(string=HTML, base_url=".").write_pdf(pdf_path)
    print(f"wrote {pdf_path}")
except Exception as e:
    print(f"PDF failed: {e}")
    pdf_path = None


# ------------------------------- DOCX -----------------------------------------
docx_path = f"{OUT_DIR}/storage_{DAY}.docx"
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def shade(cell, hexcolour):
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hexcolour.lstrip("#"))
        cell._tc.get_or_add_tcPr().append(el)

    def add_link(paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        h = OxmlElement("w:hyperlink")
        h.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        c = OxmlElement("w:color"); c.set(qn("w:val"), "0645AD"); rPr.append(c)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        run.append(rPr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        h.append(run)
        paragraph._p.append(h)

    doc = Document()
    doc.add_heading(f"Storage checks — {PRETTY}", level=1)
    doc.add_paragraph(f"{len(flagged)} of {len(jobs)} photos flagged.")

    rows_data = []
    for fid, (fname, en) in FIELDS.items():
        items = [j for j in flagged if j["field_id"] == fid]
        per = defaultdict(list)
        for j in sorted(items, key=lambda x: (x["prank"], -x["risk"])):
            per[j["code"]].append(j["ar_note"])
        for code, notes in per.items():
            note = " و".join(dict.fromkeys(x for s in notes for x in s.split(" و")))
            rows_data.append((en, note, code, COLOURS[fid][0]))

    if rows_data:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, txt in enumerate(["كود الفرع", "ملاحظات المراجعة", "مكان التخزين"]):
            hdr[i].text = txt
            shade(hdr[i], "#D9D9D9")
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for en, note, code, bg in rows_data:
            c = t.add_row().cells
            c[0].text, c[1].text, c[2].text = code, note, en
            shade(c[0], bg); shade(c[1], NOTE_BG); shade(c[2], bg)
            for cc in c:
                cc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for fid, (fname, en) in FIELDS.items():
        sub = [j for j in flagged if j["field_id"] == fid]
        if not sub:
            continue
        doc.add_heading(en, level=2)
        for j in sorted(sub, key=lambda x: (x["prank"], -x["risk"])):
            doc.add_heading(f'{j["branch"]} — {j["risk"]:.0f}%', level=3)
            doc.add_paragraph(LABELS.get(j["primary"], ""))
            try:
                im = ImageOps.exif_transpose(Image.open(j["path"])).convert("RGB")
                im.thumbnail((900, 900))
                buf = BytesIO(); im.save(buf, "JPEG", quality=80); buf.seek(0)
                doc.add_picture(buf, width=Inches(3.2))
            except Exception:
                pass
            add_link(doc.add_paragraph(), j["link"], "open submission in Zenput")

    doc.save(docx_path)
    print(f"wrote {docx_path}")
except Exception as e:
    print(f"DOCX failed: {e}")
    docx_path = None


# -------------------------------- CSV -----------------------------------------
csv_path = f"{OUT_DIR}/storage_{DAY}.csv"
with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
    cols = ["question", "branch", "code", "risk", "primary", "also", "ar_note",
            "photo_note", "flagged"] + [f"risk_{r}" for r in PRIORITY] + \
           ["user", "date", "link", "s3_key"]
    w = csv.DictWriter(f, fieldnames=cols)
    w.writeheader()
    for j in sorted(jobs, key=lambda x: (not x["flagged"], x["question"],
                                         x["prank"], -x["risk"])):
        row = {k: j.get(k, "") for k in
               ["question", "branch", "code", "primary", "ar_note", "photo_note",
                "flagged", "user", "date", "link", "s3_key"]}
        row["risk"] = round(j["risk"], 1)
        row["also"] = ", ".join(j["also"])
        for r in PRIORITY:
            row[f"risk_{r}"] = round(j["risks"][r], 1) if r in j["risks"] else ""
        w.writerow(row)
print(f"wrote {csv_path}")


# ------------------------------- EMAIL ----------------------------------------
if not SENDER_PASS:
    print("EMAIL_PASSWORD not set, skipping email")
else:
    msg = MIMEMultipart("mixed")
    msg["Subject"] = f"مراجعة تنظيم الفروع — {PRETTY}"
    msg["From"] = formataddr((SENDER_NAME, SENDER_EMAIL))
    msg["To"] = ", ".join(RECIPIENTS)
    if CC:
        msg["Cc"] = ", ".join(CC)

    body = (f'<html><body style="font:14px Tahoma,Arial">'
            f'<p dir="rtl">ملاحظات مراجعة صور تنظيم الفروع ليوم {PRETTY}:</p>'
            f'{AR_TABLE}'
            f'<p dir="rtl" style="font-size:12px;color:#777">'
            f'.</p>'
            f'</body></html>')
    msg.attach(MIMEText(body, "html", "utf-8"))

    for path in [p for p in (pdf_path, docx_path) if p and os.path.exists(p)]:
        with open(path, "rb") as f:
            part = MIMEApplication(f.read(), Name=os.path.basename(path))
        part["Content-Disposition"] = f'attachment; filename="{os.path.basename(path)}"'
        msg.attach(part)

    ctx = ssl.create_default_context()
    with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=60) as s:
        s.starttls(context=ctx)
        s.login(SENDER_EMAIL, SENDER_PASS)
        s.sendmail(SENDER_EMAIL, RECIPIENTS + CC, msg.as_string())
    print(f"emailed {len(RECIPIENTS)} recipient(s), {len(CC)} cc")

print("done")
